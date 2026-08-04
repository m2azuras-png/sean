import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Page Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        return p

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F3F4F6")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DOKUMENTASI LENGKAP & BEDAH KODE SISTEM TIMBANGAN DIGITAL ANAK")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Penjelasan Mendalam Alur Komunikasi, Baris demi Baris Kode 2 ESP32 & Web PHP Hostinger")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # BAB I: ALUR KOMUNIKASI ESP32 KE WEBSITE
    # -------------------------------------------------------------
    add_heading_1("BAB I: ALUR KOMUNIKASI ESP32 KE WEBSITE")
    
    p = doc.add_paragraph()
    p.add_run("Sistem ini bekerja dengan memadukan dua mikrokontroler ESP32 dan satu server web publik (Hostinger). Komunikasi terbagi menjadi dua ranah: ranah lokal nirkabel (ESP-NOW) dan ranah internet publik (HTTPS & REST API).")

    add_heading_2("1.1 Tahapan Alur Komunikasi")
    p = doc.add_paragraph()
    p.add_run("1. Pengukuran Berat Lokal (ESP32 Bawah -> ESP32 Utama):\n").bold = True
    p.add_run("   ESP32 Bawah membaca sinyal analog dari sensor Loadcell HX711 di pijakan kaki. Data berat dikonversi ke satuan Kilogram (kg) lalu dipancarkan via frekuensi radio 2.4GHz tanpa router (ESP-NOW Broadcast) setiap 100 milidetik.\n\n")
    p.add_run("2. Pengolahan Data Multi-Sensor (ESP32 Utama):\n").bold = True
    p.add_run("   ESP32 Utama bertindak sebagai titik kumpul. Ketika data berat diterima via callback OnDataRecv, ESP32 Utama secara bersamaan membaca sensor Ultrasonik JSN-SR04T (tinggi badan) dan Optocoupler (lingkar perut), lalu menampilkannya pada layar LCD 20x4.\n\n")
    p.add_run("3. Pengiriman HTTPS POST (ESP32 Utama -> Server Hostinger):\n").bold = True
    p.add_run("   Saat tombol fisik 'K' ditekan, ESP32 Utama membuka socket SSL aman (WiFiClientSecure) dan mengirimkan payload JSON berisi API Key dan nilai sensor ke https://seanta.my.id/api_sensor.php.\n\n")
    p.add_run("4. Penyimpanan & Tampilan Real-Time (Server -> Browser):\n").bold = True
    p.add_run("   File api_sensor.php melakukan verifikasi API Key dan menyimpan data ke database MySQL (bacaan_sensor). Di saat bersamaan, halaman browser petugas (pengukuran.php) secara terus-menerus mengambil data terbaru (polling 1 detik) dari baca_sensor.php dan mengisi nilai form secara otomatis tanpa refresh.")

    # -------------------------------------------------------------
    # BAB II: BEDAH KODE PROGRAM ESP32 (HARDWARE)
    # -------------------------------------------------------------
    add_heading_1("BAB II: BEDAH KODE PROGRAM ESP32 (HARDWARE)")

    add_heading_2("2.1 Kode ESP32 Bawah / Transmitter (espbawah.ino)")
    p = doc.add_paragraph()
    p.add_run("ESP32 ini khusus menangani timbangan berat badan menggunakan Loadcell HX711.")

    add_code_block('''#include <HX711_ADC.h>
#include <WiFi.h>
#include <esp_now.h>

const int HX711_dout = 25;
const int HX711_sck  = 26;
HX711_ADC LoadCell(HX711_dout, HX711_sck);
float calibrationValue = 50194.0;

const char* WIFI_SSID     = "PERTAMINA123_plus";
const char* WIFI_PASSWORD = "Bismillah";

typedef struct struct_message {
  float berat;
} struct_message;

struct_message myData;
uint8_t broadcastAddress[] = {0xFF, 0xFF, 0xFF, 0xFF, 0xFF, 0xFF};
esp_now_peer_info_t peerInfo;

unsigned long lastTime = 0;
const unsigned long timerDelay = 100;

void setup() {
  Serial.begin(115200);
  WiFi.mode(WIFI_STA);
  WiFi.begin(WIFI_SSID, WIFI_PASSWORD);
  
  while (WiFi.status() != WL_CONNECTED) { delay(500); }

  if (esp_now_init() != ESP_OK) return;

  memcpy(peerInfo.peer_addr, broadcastAddress, 6);
  peerInfo.channel = 0;  
  peerInfo.encrypt = false;
  esp_now_add_peer(&peerInfo);

  LoadCell.begin();
  LoadCell.start(2000, true);
  LoadCell.setCalFactor(calibrationValue);
}

void loop() {
  LoadCell.update();
  if (millis() - lastTime > timerDelay) {
    myData.berat = LoadCell.getData();
    esp_now_send(broadcastAddress, (uint8_t *) &myData, sizeof(myData));
    lastTime = millis();
  }
}''')

    add_heading_3("Penjelasan Detail Kode espbawah.ino:")
    p = doc.add_paragraph()
    p.add_run("• #include <HX711_ADC.h>: Library pengolah sinyal analog-to-digital converter dari penguat HX711 Loadcell.\n")
    p.add_run("• HX711_dout = 25 & HX711_sck = 26: Mendefinisikan Pin GPIO 25 sebagai jalur data dan Pin GPIO 26 sebagai jalur clock komunikasi HX711.\n")
    p.add_run("• calibrationValue = 50194.0: Nilai pembagi kalibrasi fisik agar tegangan regangan strain gauge terkonversi tepat ke satuan Kg.\n")
    p.add_run("• struct struct_message { float berat; }: Struktur data berisi variabel float yang disepakati bersama antara transmitter dan receiver ESP-NOW.\n")
    p.add_run("• broadcastAddress[] = {0xFF,0xFF,0xFF,0xFF,0xFF,0xFF}: Menggunakan alamat MAC broadcast universal agar ESP32 bawah dapat memancarkan data tanpa perlu mengetahui MAC address fisik ESP32 utama.\n")
    p.add_run("• WiFi.mode(WIFI_STA) & WiFi.begin(): Mengaktifkan radio WiFi ESP32 untuk menyamakan channel radio dengan akses poin WiFi agar transmisi ESP-NOW lancar.\n")
    p.add_run("• LoadCell.update(): Fungsi non-blocking yang wajib dipanggil di loop untuk terus mengambil sampel data berat terkini.\n")
    p.add_run("• esp_now_send(): Mengirimkan paket data myData berukuran sizeof(myData) secara nirkabel setiap 100ms.")

    add_heading_2("2.2 Kode ESP32 Utama / Main Controller (esp32_timbangan.ino)")
    p = doc.add_paragraph()
    p.add_run("ESP32 Utama merupakan pusat pengendali yang mengintegrasikan ESP-NOW, Sensor Ultrasonik, Optocoupler, LCD I2C 20x4, Pemutar Audio DFPlayer Mini, dan Client HTTPS POST.")

    add_heading_3("Penjelasan Fungsi-Fungsi Kunci esp32_timbangan.ino:")
    p = doc.add_paragraph()
    p.add_run("1. Callback OnDataRecv():\n").bold = True
    p.add_run("   Fungsi otomatis yang dipanggil oleh sistem operasi ESP32 saat ada paket ESP-NOW yang masuk. Data disalin ke variabel beratMentahReceiver via memcpy().\n\n")
    p.add_run("2. Fungsi bacaTinggi():\n").bold = True
    p.add_run("   Memancar pulsa 10 mikrodetik pada TRIG_PIN (18) dan mengukur durasi pantulan pada ECHO_PIN (19). Jarak dihitung dengan rumus d = durasi * 0.0343 / 2.0. Fungsi ini dilengkapi pemfilteran loncatan drastis (>30cm) dan penyeimbang eksponensial (0.7 * lama + 0.3 * baru) agar angka tinggi badan tidak bergetar.\n\n")
    p.add_run("3. Fungsi updatePulsaKeliling():\n").bold = True
    p.add_run("   Membaca perubahan sinyal HIGH/LOW dari Pin 2 optocoupler. Menggunakan algoritma debounce (15ms) untuk mencegah penghitungan pulsa ganda saat pita ditarik.\n\n")
    p.add_run("4. Fungsi kirimDataKeServer() (HTTPS Anti-Freeze):\n").bold = True
    p.add_run("   - Membuat objek WiFiClientSecure client dan memanggil client.setInsecure() untuk melewati validasi sertifikat SSL Hostinger.\n")
    p.add_run("   - Menyetel timeout client.setTimeout(5000) dan http.setTimeout(5000) agar ESP32 tidak freeze jika koneksi lambat.\n")
    p.add_run("   - Menyusun string JSON payload dari variabel beratHalus, tinggi_cm, dan keliling_cm.\n")
    p.add_run("   - Mengirim request POST via http.POST(payload) dan wajib mengakhiri dengan client.stop() untuk menghapus memori SSL RAM agar ESP32 tidak kehabisan memori.\n\n")
    p.add_run("5. Fungsi aksiKirimData():\n").bold = True
    p.add_run("   Memaksa pembacaan sensor tinggi secara instan tepat sebelum dikirim, memanggil kirimDataKeServer(), lalu membunyikan file audio MP3 folder 4 pada DFPlayer Mini.")

    # -------------------------------------------------------------
    # BAB III: BEDAH KODE WEBSITE (PHP NATIVE & MYSQL)
    # -------------------------------------------------------------
    add_heading_1("BAB III: BEDAH KODE WEBSITE (PHP NATIVE & MYSQL)")

    add_heading_2("3.1 config.php")
    p = doc.add_paragraph()
    p.add_run("• $koneksi = new mysqli(...): Membuat koneksi MySQL ke database Hostinger (u114571555_sean).\n")
    p.add_run("• define('ESP32_API_KEY', 'Tmbgn-Pertamina-XYZ998'): Konstanta keamanan kunci rahasia yang wajib sama persis dengan yang dikirim oleh ESP32.\n")
    p.add_run("• hitung_status_gizi($usia, $bmi): Algoritma penentu kategori gizi (Sangat Kurus, Kurus, Normal, Gemuk, Obesitas) berdasarkan indeks massa tubuh (BMI = BB / (TB_m)^2) dan rentang usia.")

    add_heading_2("3.2 api_sensor.php (Endpoint Receiver ESP32)")
    p = doc.add_paragraph()
    p.add_run("• json_decode(file_get_contents('php://input'), true): Membaca body request bertipe JSON yang dikirimkan oleh ESP32 via HTTP POST.\n")
    p.add_run("• Validasi API Key: Mengecek apakah ($input['api_key'] === ESP32_API_KEY). Jika beda, mengembalikan HTTP Response Code 401 (Unauthorized).\n")
    p.add_run("• SQL UPSERT: INSERT INTO bacaan_sensor (id, berat_badan, tinggi_badan, lingkar_perut) VALUES (1, ?, ?, ?) ON DUPLICATE KEY UPDATE ...: Menjamin data baris id=1 selalu ada dan langsung terbarui tanpa membuat baris baru yang menumpuk.")

    add_heading_2("3.3 baca_sensor.php (Endpoint Polling Frontend)")
    p = doc.add_paragraph()
    p.add_run("• Header Cache-Control: no-store, no-cache: Memerintahkan browser pengakses untuk TIDAK menyimpan cache respons HTTP. Setiap kali dipanggil, browser dipaksa mengambil data yang paling fresh dari MySQL Hostinger.\n")
    p.add_run("• $data['detik_lalu'] = time() - (int)$data['last_update']: Menghitung selisih detik antara jam server sekarang dengan waktu data sensor terakhir masuk, untuk ditampilkan pada indikator status alat di web.")

    add_heading_2("3.4 pengukuran.php (Polling JavaScript Real-Time)")
    p = doc.add_paragraph()
    p.add_run("• fetch('baca_sensor.php?_t=' + Date.now(), { cache: 'no-store' }): Fungsi JavaScript yang dipanggil setiap 1000ms (1 detik) untuk mengambil JSON dari server.\n")
    p.add_run("• document.activeElement !== el: Memastikan nilai pada kotak input (berat, tinggi, lingkar) terisi otomatis dari sensor, TANPA mengganggu/menimpa jika petugas sedang mengetik manual di kolom tersebut.\n")
    p.add_run("• Indikator ● Alat Terhubung: Berubah warna hijau secara otomatis jika detik_lalu <= 15 detik, menandakan ESP32 aktif mengirimkan data.")

    # Save document
    doc_path = "c:\\xampp\\htdocs\\Timbangan\\Dokumentasi_Sistem_Timbangan_Digital.docx"
    doc.save(doc_path)
    print("Dokumen Word berhasil diperbarui:", doc_path)

if __name__ == "__main__":
    create_document()
